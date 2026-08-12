import streamlit as st
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# =========================================================
# PAGE
# =========================================================

st.set_page_config(
    page_title="AI Resume Builder",
    page_icon="📄",
    layout="wide"
)

st.title("📄 AI Resume Builder")
st.caption("Professional job-targeted resume builder • No API key required")

# =========================================================
# RESET
# =========================================================

if st.button("🔄 Start New Resume"):
    st.session_state.clear()
    st.rerun()

# =========================================================
# TEMPLATE
# =========================================================

template = st.selectbox(
    "🎨 Choose Resume Template",
    ["Simple", "Professional", "Modern"]
)

# =========================================================
# JOB
# =========================================================

st.header("🎯 Target Job")

job_title = st.text_input(
    "Target Job Title",
    placeholder="Example: Python Developer"
)

job_description = st.text_area(
    "Job Description",
    height=160,
    placeholder="Paste the job description here..."
)

# =========================================================
# PERSONAL
# =========================================================

st.header("👤 Personal Information")

profile_photo = st.file_uploader(
    "📷 Profile Photo",
    type=["jpg", "jpeg", "png"]
)

name = st.text_input("Full Name")
email = st.text_input("Email")
phone = st.text_input("Phone")
location = st.text_input("Location")

# =========================================================
# EDUCATION
# =========================================================

st.header("🎓 Education")

education = st.text_area(
    "Education",
    placeholder="BSc Computer Science - College Name - 2026"
)

# =========================================================
# SKILLS
# =========================================================

st.header("🛠️ Skills")

skills = st.text_area(
    "Skills",
    placeholder="Python, SQL, HTML, CSS, JavaScript, Git"
)

# =========================================================
# PROJECTS
# =========================================================

st.header("🚀 Projects")

projects = st.text_area(
    "Projects",
    placeholder="AI Resume Builder using Python and Streamlit"
)

# =========================================================
# EXPERIENCE
# =========================================================

st.header("💼 Experience")

experience = st.text_area(
    "Experience",
    placeholder="Fresher"
)

# =========================================================
# CERTIFICATIONS
# =========================================================

st.header("🏆 Certifications")

certifications = st.text_area(
    "Certifications",
    placeholder="Python Full Stack Certification"
)

# =========================================================
# ACHIEVEMENTS
# =========================================================

st.header("⭐ Achievements")

achievements = st.text_area(
    "Achievements",
    placeholder="Completed multiple Python projects"
)

# =========================================================
# SUMMARY
# =========================================================

st.header("✨ Professional Summary")

if st.button("✨ Generate Professional Summary"):

    if not job_title:
        st.warning("Enter Target Job Title.")

    elif not education:
        st.warning("Enter Education.")

    elif not skills:
        st.warning("Enter Skills.")

    else:

        summary = (
            f"Motivated {education} graduate seeking an "
            f"entry-level {job_title} position. "
            f"Skilled in {skills}. "
            f"Developed projects including {projects}. "
            f"Eager to apply technical knowledge, "
            f"learn new technologies, solve problems, "
            f"and contribute effectively to a professional team."
        )

        st.session_state["summary"] = summary

        st.success("Professional summary generated!")

summary = st.session_state.get("summary", "")

if summary:

    st.text_area(
        "✏️ Edit Professional Summary",
        value=summary,
        key="summary_editor",
        height=130
    )

    summary = st.session_state["summary_editor"]

# =========================================================
# KEYWORD FUNCTION
# =========================================================

def extract_keywords(text):

    words = re.findall(
        r"[A-Za-z][A-Za-z+#.-]{2,}",
        text.lower()
    )

    stop_words = {
        "the", "and", "for", "with", "you",
        "are", "our", "this", "that", "from",
        "will", "have", "has", "job", "work",
        "using", "looking", "years", "role",
        "candidate", "into", "your", "their",
        "they", "who", "all", "can", "should"
    }

    return set(
        word for word in words
        if word not in stop_words
    )

# =========================================================
# ATS KEYWORD MATCH
# =========================================================

if st.button("🔎 Check ATS Keyword Match"):

    if not job_description:
        st.warning("Paste a job description first.")

    elif not skills:
        st.warning("Enter your skills first.")

    else:

        job_words = extract_keywords(job_description)
        skill_words = extract_keywords(
            skills + " " + projects + " " + experience
        )

        matching = sorted(
            job_words.intersection(skill_words)
        )

        missing = sorted(
            job_words.difference(skill_words)
        )

        if job_words:
            match = int(
                len(matching) / len(job_words) * 100
            )
        else:
            match = 0

        st.subheader("🔎 ATS Keyword Match")

        st.progress(match / 100)

        st.write(f"## {match}%")

        col1, col2 = st.columns(2)

        with col1:

            st.success("✅ Matching Keywords")

            if matching:
                st.write(", ".join(matching))
            else:
                st.write("No matching keywords found.")

        with col2:

            st.warning("⚠️ Possible Missing Keywords")

            if missing:
                st.write(", ".join(missing[:25]))
            else:
                st.write("No obvious missing keywords.")

# =========================================================
# RESUME SCORE
# =========================================================

def calculate_score():

    score = 0

    if name:
        score += 10

    if email:
        score += 10

    if phone:
        score += 5

    if location:
        score += 5

    if job_title:
        score += 10

    if summary:
        score += 10

    if education:
        score += 10

    if skills:
        score += 10

    if projects:
        score += 10

    if experience:
        score += 5

    if certifications:
        score += 5

    if achievements:
        score += 5

    return score

if st.button("📊 Check Resume Score"):

    score = calculate_score()

    st.subheader("📊 Resume Score")

    st.progress(score / 100)

    st.write(f"## {score}%")

    if score >= 80:
        st.success("Excellent resume!")

    elif score >= 60:
        st.info("Good resume. Add more details if possible.")

    elif score >= 40:
        st.warning("Your resume needs more information.")

    else:
        st.error("Complete more sections.")

# =========================================================
# RESUME TEXT
# =========================================================

def resume_text():

    text = f"""
{name}

{job_title}

{location} | {email} | {phone}

PROFESSIONAL SUMMARY

{summary}

SKILLS

{skills}

EDUCATION

{education}

PROJECTS

{projects}

EXPERIENCE

{experience}

CERTIFICATIONS

{certifications}

ACHIEVEMENTS

{achievements}
"""

    return text.strip()

# =========================================================
# COVER LETTER
# =========================================================

st.header("📝 Cover Letter")

if st.button("📝 Generate Cover Letter"):

    if not name:
        st.warning("Enter your name.")

    elif not job_title:
        st.warning("Enter Target Job Title.")

    else:

        cover_letter = f"""Dear Hiring Manager,

I am writing to express my interest in the {job_title} position.

I am a {education} graduate with skills in {skills}. I have worked on projects such as {projects}.

My technical knowledge, willingness to learn, and problem-solving mindset would allow me to contribute positively to your organization.

I would appreciate the opportunity to discuss how my skills and background can contribute to your team.

Thank you for considering my application.

Sincerely,
{name}
{email}
{phone}
"""

        st.session_state["cover_letter"] = cover_letter

        st.success("Cover letter generated!")

if "cover_letter" in st.session_state:

    st.text_area(
        "✏️ Edit Cover Letter",
        value=st.session_state["cover_letter"],
        height=300,
        key="cover_editor"
    )

# =========================================================
# GENERATE RESUME
# =========================================================

if st.button("📄 Generate Resume"):

    if not name:
        st.warning("Please enter your name.")

    elif not email:
        st.warning("Please enter your email.")

    elif not job_title:
        st.warning("Please enter Target Job.")

    else:

        st.success("Resume generated successfully!")

        # =================================================
        # PREVIEW
        # =================================================

        st.header("📋 Resume Preview")

        if profile_photo:

            preview = Image.open(profile_photo)

            st.image(
                preview,
                width=150
            )

        st.subheader(name)

        st.write(f"🎯 {job_title}")

        st.write(
            f"{location} | {email} | {phone}"
        )

        st.markdown("### PROFESSIONAL SUMMARY")
        st.write(summary)

        st.markdown("### SKILLS")
        st.write(skills)

        st.markdown("### EDUCATION")
        st.write(education)

        st.markdown("### PROJECTS")
        st.write(projects)

        st.markdown("### EXPERIENCE")
        st.write(experience)

        if certifications:
            st.markdown("### CERTIFICATIONS")
            st.write(certifications)

        if achievements:
            st.markdown("### ACHIEVEMENTS")
            st.write(achievements)

        # =================================================
        # PDF
        # =================================================

        pdf = BytesIO()

        c = canvas.Canvas(
            pdf,
            pagesize=A4
        )

        width, height = A4

        y = height - 50

        if template == "Simple":

            name_font = 20
            heading_font = 12

        elif template == "Professional":

            name_font = 24
            heading_font = 14

        else:

            name_font = 26
            heading_font = 15

        # PHOTO

        if profile_photo:

            profile_photo.seek(0)

            photo = Image.open(
                profile_photo
            )

            buffer = BytesIO()

            photo.save(
                buffer,
                format="PNG"
            )

            buffer.seek(0)

            c.drawImage(
                ImageReader(buffer),
                width - 140,
                height - 140,
                width=80,
                height=80,
                preserveAspectRatio=True,
                mask="auto"
            )

        # NAME

        c.setFont(
            "Helvetica-Bold",
            name_font
        )

        c.drawString(
            50,
            y,
            name
        )

        y -= 25

        # JOB

        c.setFont(
            "Helvetica-Bold",
            11
        )

        c.drawString(
            50,
            y,
            job_title
        )

        y -= 20

        # CONTACT

        c.setFont(
            "Helvetica",
            9
        )

        c.drawString(
            50,
            y,
            f"{location} | {email} | {phone}"[:110]
        )

        y -= 35

        # SUMMARY

        if summary:

            c.setFont(
                "Helvetica-Bold",
                heading_font
            )

            c.drawString(
                50,
                y,
                "PROFESSIONAL SUMMARY"
            )

            y -= 18

            c.setFont(
                "Helvetica",
                9
            )

            words = summary.split()

            line = ""

            for word in words:

                if len(line + " " + word) < 95:

                    line += " " + word

                else:

                    c.drawString(
                        50,
                        y,
                        line.strip()
                    )

                    y -= 13

                    line = word

            if line:

                c.drawString(
                    50,
                    y,
                    line.strip()
                )

                y -= 25

        # SECTIONS

        sections = [
            ("SKILLS", skills),
            ("EDUCATION", education),
            ("PROJECTS", projects),
            ("EXPERIENCE", experience),
            ("CERTIFICATIONS", certifications),
            ("ACHIEVEMENTS", achievements)
        ]

        for title, content in sections:

            if content:

                c.setFont(
                    "Helvetica-Bold",
                    heading_font
                )

                c.drawString(
                    50,
                    y,
                    title
                )

                y -= 18

                c.setFont(
                    "Helvetica",
                    9
                )

                for line in content.split("\n"):

                    c.drawString(
                        50,
                        y,
                        line[:110]
                    )

                    y -= 13

                y -= 12

                # New page if needed

                if y < 60:

                    c.showPage()

                    y = height - 50

        c.save()

        pdf.seek(0)

        st.download_button(
            "📥 Download Resume PDF",
            data=pdf,
            file_name=f"{job_title.replace(' ', '_')}_Resume.pdf",
            mime="application/pdf"
        )

        # =================================================
        # DOCX
        # =================================================

        document = Document()

        # PHOTO

        if profile_photo:

            profile_photo.seek(0)

            image_buffer = BytesIO(
                profile_photo.read()
            )

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = paragraph.add_run()

            run.add_picture(
                image_buffer,
                width=Inches(1)
            )

        # NAME

        p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(name)

        run.bold = True

        run.font.size = Pt(22)

        # JOB

        p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(job_title)

        run.bold = True

        # CONTACT

        p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run(
            f"{location} | {email} | {phone}"
        )

        # DOCX SECTIONS

        doc_sections = [
            ("PROFESSIONAL SUMMARY", summary),
            ("SKILLS", skills),
            ("EDUCATION", education),
            ("PROJECTS", projects),
            ("EXPERIENCE", experience),
            ("CERTIFICATIONS", certifications),
            ("ACHIEVEMENTS", achievements)
        ]

        for title, content in doc_sections:

            if content:

                heading = document.add_heading(
                    title,
                    level=2
                )

                paragraph = document.add_paragraph(
                    content
                )

        docx_file = BytesIO()

        document.save(docx_file)

        docx_file.seek(0)

        st.download_button(
            "📝 Download Resume DOCX",
            data=docx_file,
            file_name=f"{job_title.replace(' ', '_')}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # =================================================
        # COPYABLE TEXT
        # =================================================

        st.subheader("📋 Resume Text")

        st.text_area(
            "Copy your resume text",
            value=resume_text(),
            height=300
        )

# =========================================================
# FOOTER
# =========================================================

st.divider()

st.caption(
    "AI Resume Builder | Python + Streamlit | No API Key Required"
)